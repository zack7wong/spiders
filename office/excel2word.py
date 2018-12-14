#!/usr/bin/env python
# -*- coding:utf-8 -*-

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pyexcel_xls import get_data
from docx.shared import Pt
from docx import Document
from docx.oxml.ns import qn
import time


def write(res):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #大标题居中
    title = document.add_paragraph('个人住房按揭贷款业务调查报告')
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style = document.styles.add_style('title' ,WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(18)
    style.font.bold = True
    title.style = style

    #空一行
    # document.add_paragraph('')

    #开始
    start = document.add_paragraph('一、借款申请人基本情况')
    paragraph_format = start.paragraph_format
    paragraph_format.space_after = Pt(0)
    style = document.styles.add_style('start' ,WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(12)
    style.font.bold = True
    start.style = style

    #第一段
    one = document.add_paragraph('    借款申请人')
    paragraph_format = one.paragraph_format
    paragraph_format.space_after = Pt(0)
    one.add_run(' '+res['name']+' ').underline = True
    one.add_run('，性别')
    one.add_run(' '+res['sex']+' ').underline = True
    one.add_run('，年龄')
    #计算年龄
    nowyear = time.strftime('%Y',time.localtime(time.time()))
    age = str(int(nowyear) - int(res['idcard'][6:10]))

    one.add_run(' '+age+' ').underline = True
    one.add_run('岁，联系电话')
    one.add_run(' '+res['phone']+' ').underline = True
    one.add_run('，身份证号码')
    one.add_run(' '+res['idcard']+' ').underline = True
    one.add_run('，户籍所在地')
    one.add_run(' '+res['address1']+' ').underline = True
    one.add_run('，现住址')
    one.add_run(' '+res['address2']+' ').underline = True
    one.add_run('，工作单位')
    one.add_run(' '+res['work']+' ').underline = True
    one.add_run('，婚姻状况为')
    one.add_run(' 已婚 ').underline = True
    one.add_run('，家庭供养人口')
    one.add_run(' '+res['people']+' ').underline = True
    one.add_run('人，经查询借款人个人征信')
    one.add_run(' '+res['yuqi']+' ').underline = True
    one.add_run('逾期记录。')

    #第二段
    one = document.add_paragraph('    财产共有人（配偶）')
    paragraph_format = one.paragraph_format
    paragraph_format.space_after = Pt(0)
    one.add_run(' '+res['peiou']+' ').underline = True
    one.add_run('，身份证号码')
    one.add_run(' '+res['peiou_idcard']+' ').underline = True
    one.add_run('，工作单位')
    one.add_run(' '+res['peiou_work']+'。 ').underline = True


    #第三段
    one = document.add_paragraph('    本次借款用途为')
    paragraph_format = one.paragraph_format
    paragraph_format.space_after = Pt(10)
    one.add_run(' 个人住房按揭 ').underline = True
    one.add_run('，现已支付首期款￥')
    one.add_run(' '+res['shoufu']+' ').underline = True
    one.add_run('元，占所购买房屋总价值￥')
    one.add_run(' '+res['allprice']+' ').underline = True
    one.add_run('元的')

    shoufu_pre = (float(res['shoufu']) / float(res['allprice']))*100
    shoufu_pre = str("%.2f" % shoufu_pre)
    one.add_run(' '+shoufu_pre+'% ').underline = True
    one.add_run('。申请个人贷款金额为￥')

    daikuan = str(int(res['allprice']) - int(res['shoufu']))
    one.add_run(' '+daikuan+' ').underline = True
    one.add_run('元，期限')
    one.add_run(' '+res['daikuanqixian']+' ').underline = True
    one.add_run('年，贷款额度占所抵押资产总价的')

    edu = (float(daikuan) / float(res['allprice']))*100
    edu = str("%.2f" % edu)
    one.add_run(' '+edu+'%。 ').underline = True

    #中间
    zhongjian = document.add_paragraph('二、借款申请人还款保障状况')
    paragraph_format = zhongjian.paragraph_format
    paragraph_format.space_after = Pt(0)
    style = document.styles.add_style('zhongjian' ,WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(12)
    style.font.bold = True
    zhongjian.style = style

    #第四段
    one = document.add_paragraph('    根据借款申请人提供的')
    paragraph_format = one.paragraph_format
    paragraph_format.space_after = Pt(0)
    one.add_run(' 荆门市商品房买卖 ').underline = True
    one.add_run('合同及其首期付款收据，申请人及配偶身份证明、经济收入证明，财产共有人出具的承诺书等资料，经本人实地调查核实均真实、合法、有效。')

    #第五段
    one = document.add_paragraph('    1、申请人及其配偶（家庭）月收入为￥')
    paragraph_format = one.paragraph_format
    paragraph_format.space_after = Pt(0)
    one.add_run(' '+res['yueshouru']+' ').underline = True
    one.add_run('元，而借款申请人本次贷款月还款额为')
    one.add_run(' '+res['yuehuankuan']+' ').underline = True
    one.add_run('元，月物业管理费为')
    one.add_run(' 0 ').underline = True
    one.add_run('元，占月均收入的')

    yuejun = (float(res['yuehuankuan']) / float(res['yueshouru']))*100
    yuejun = str("%.2f" % yuejun)
    one.add_run(' '+yuejun+'% ').underline = True

    one.add_run('。申请人及其配偶（家庭）其他债务月均偿付额为')
    one.add_run(' '+res['jiating']+' ').underline = True
    one.add_run('元，所有债务合计')

    zhaiwu = float(res['yuehuankuan']) + float(res['jiating'])
    zhaiwu = str("%.2f" % zhaiwu)
    one.add_run(' '+zhaiwu+' ').underline = True

    one.add_run('元（本次贷款月还款额+月物业管理费+其他债务月均偿付额）占月均收入的')
    zhaiwu_pre = (float(zhaiwu) / float(res['yueshouru']))*100
    zhaiwu_pre = str("%.2f" % zhaiwu_pre)
    one.add_run(' '+zhaiwu_pre+'% ').underline = True
    one.add_run('。借款申请人经济状况较好，收入较稳定，第一还款来源充足。')

    #第六段
    one = document.add_paragraph('    2、借款申请人以')
    paragraph_format = one.paragraph_format
    paragraph_format.space_after = Pt(10)
    one.add_run(' 房地产 ').underline = True
    one.add_run('作为借款的 抵押担保，第二还款来源充足，有关手续合法有效。抵押物为')
    one.add_run(' 房地产 ').underline = True
    one.add_run('，位于')
    one.add_run(' '+res['fangchan']+' ').underline = True
    one.add_run('，价值为')
    one.add_run(' '+res['allprice']+' ').underline = True
    one.add_run('元。')

    #中间
    end = document.add_paragraph('三、调查意见')
    paragraph_format = end.paragraph_format
    paragraph_format.space_after = Pt(0)
    style = document.styles.add_style('end' ,WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(12)
    style.font.bold = True
    end.style = style

    #第七段
    one = document.add_paragraph('    通过对借款申请人的调查核实，借款申请人第一和第二还款来源均有保障，符合个人贷款条件，本人同意对借款人')
    paragraph_format = one.paragraph_format
    paragraph_format.space_after = Pt(10)
    one.add_run(' '+res['name']+' ').underline = True
    one.add_run('发放个人住房按揭贷款￥')
    daikuan_price = str(int(res['allprice'])- int(res['shoufu']))
    one.add_run(' '+daikuan_price+' ').underline = True
    one.add_run('元，期限')
    one.add_run(' '+res['daikuanqixian']+' ').underline = True
    one.add_run('年，利率')
    one.add_run(' 5.88% ').underline = True
    one.add_run('。同时在贷款审批完毕后，本人将及时办理抵押登记等相关手续，从而全面防范贷款风险。')

    #第八段
    one = document.add_paragraph('                                                    调查人：')
    #第九段
    # one = document.add_paragraph('')
    #第十段
    one = document.add_paragraph('                                                    年  月  日')

    filename = res['name']+'.docx'
    document.save(filename)

def read_excel(filename):
    xls_data = get_data(filename)
    results = []
    for sheet_n in xls_data.keys():
        for item in xls_data[sheet_n][1:]:
            if item[0]=='姓名' and item[1] == '性别':
                continue
            print(item)
            name = item[0]
            sex = item[1]
            phone = item[2]
            idcard = item[3]
            address1 = item[4]
            address2 = item[5]
            work = item[6]
            people = item[7]
            yuqi = item[8]
            peiou = item[9]
            peiou_idcard = item[10]
            peiou_work = item[11]
            shoufu = item[12]
            allprice = item[13]
            daikuanqixian = item[14]
            yueshouru = item[15]
            yuehuankuan = item[16]
            jiating = item[17]
            fangchan = item[18]

            obj ={
                'name':name,
                'sex':sex,
                'phone':phone,
                'idcard':idcard,
                'address1':address1,
                'address2':address2,
                'work':work,
                'people':people,
                'yuqi':yuqi,
                'peiou':peiou,
                'peiou_idcard':peiou_idcard,
                'peiou_work':peiou_work,
                'shoufu':shoufu,
                'allprice':allprice,
                'daikuanqixian':daikuanqixian,
                'yueshouru':yueshouru,
                'yuehuankuan':yuehuankuan,
                'jiating':jiating,
                'fangchan': fangchan
            }
            results.append(obj)

    return results

def main():
    results = read_excel('客户信息.xls')
    for res in results:
        write(res)

if __name__ == '__main__':
    main()